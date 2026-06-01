from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document('sprawozdanie.docx')

# ── 1. Zmień kolory nagłówków ──────────────────────────────────────────────
COLOR_H1 = RGBColor(0x1a, 0x1a, 0x2e)  # granat (kolor z aplikacji)
COLOR_H2 = RGBColor(0xe9, 0x45, 0x60)  # czerwień (kolor z aplikacji)

for style in doc.styles:
    if style.name == 'Heading 1':
        style.font.color.rgb = COLOR_H1
        style.font.bold = True
        style.font.size = Pt(14)
        style.font.underline = False
    elif style.name == 'Heading 2':
        style.font.color.rgb = COLOR_H2
        style.font.bold = True
        style.font.size = Pt(12)
        style.font.underline = False

# Wymuś kolor bezpośrednio na runach — nadpisuje formatowanie inline z AI
for para in doc.paragraphs:
    if para.style.name == 'Heading 1':
        for run in para.runs:
            run.font.color.rgb = COLOR_H1
    elif para.style.name == 'Heading 2':
        for run in para.runs:
            run.font.color.rgb = COLOR_H2

# ── 2. Poprawki treści ─────────────────────────────────────────────────────
paras_to_remove = []

for para in doc.paragraphs:
    text = para.text

    # 8 tabel → 6 tabel
    if '8 tabel' in text:
        for run in para.runs:
            run.text = run.text.replace('8 tabel', '6 tabel')

    # Usuń błędną wzmiankę o ON DELETE CASCADE (nie mamy tego w schemacie)
    if 'ON DELETE CASCADE' in text:
        for run in para.runs:
            run.text = run.text.replace(
                '; klucze obce z ON DELETE CASCADE tam, gdzie logicznie uzasadnione',
                ' oraz klucze obce zapewniające integralność referencyjną'
            )

    # Oznacz do usunięcia akapity o private_leagues / private_league_members
    if 'private_leagues' in text or 'private_league_members' in text:
        paras_to_remove.append(para)

for para in paras_to_remove:
    para._element.getparent().remove(para._element)

# ── 3. Przenumeruj sekcje: 11→12, 10.1→11.1, 10→11 ───────────────────────
# Kolejność ważna — najpierw wyższe numery żeby nie trafić dwa razy
for para in doc.paragraphs:
    if para.text.strip() == '11. Wnioski':
        for run in para.runs:
            run.text = run.text.replace('11.', '12.', 1)
        break

for para in doc.paragraphs:
    if para.text.strip().startswith('10.1.'):
        for run in para.runs:
            run.text = run.text.replace('10.1.', '11.1.', 1)

for para in doc.paragraphs:
    if para.text.strip() == '10. Testowanie':
        for run in para.runs:
            run.text = run.text.replace('10.', '11.', 1)
        break

# ── 4. Wstaw sekcję 10: Microsoft Access ──────────────────────────────────
def get_style_id(doc, style_name):
    for style in doc.styles:
        if style.name == style_name:
            return style.element.get(qn('w:styleId'))
    return 'Normal'

def make_para(doc, text, style_name):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), get_style_id(doc, style_name))
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

# Znajdź "11. Testowanie" (po przenumerowaniu) — przed nim wstawiamy sekcję 10
insert_before = None
for para in doc.paragraphs:
    if para.text.strip() == '11. Testowanie':
        insert_before = para
        break

if insert_before:
    access_content = [
        ('10. Odwzorowanie schematu w Microsoft Access', 'Heading 1'),
        (
            'Schemat bazy danych został odwzorowany w Microsoft Access jako niezależna kopia struktury. '
            'Plik db/access_schema.sql zawiera definicje wszystkich 6 tabel w dialekcie Jet SQL (Access SQL). '
            'Tabele odwzorowują ten sam schemat co baza PostgreSQL — różnią się wyłącznie składnią.',
            'Normal'
        ),
        ('Kluczowe różnice między dialektem PostgreSQL a Access SQL:', 'Normal'),
        ('SERIAL → COUNTER (automatyczny licznik klucza głównego)', 'List Bullet'),
        ('VARCHAR(n) → TEXT(n)', 'List Bullet'),
        ('SMALLINT → SHORT', 'List Bullet'),
        ('TIMESTAMPTZ → DATETIME', 'List Bullet'),
        ("DEFAULT NOW() → DEFAULT Now()", 'List Bullet'),
        (
            "Ograniczenia CHECK (np. status IN ('scheduled', 'finished')) nie są wspierane "
            'w SQL Access — ustawiane ręcznie przez właściwość Validation Rule w widoku projektu tabeli.',
            'List Bullet'
        ),
        ('10.1. Kolejność tworzenia tabel', 'Heading 2'),
        (
            'Ze względu na klucze obce tabele muszą być tworzone w określonej kolejności: '
            'users, leagues, scoring_rules, następnie teams, matches, a na końcu predictions. '
            'Każdą instrukcję CREATE TABLE należy wykonać osobno w edytorze SQL Accessa '
            '(Create → Query Design → SQL View → Run).',
            'Normal'
        ),
    ]

    ref = insert_before._element
    for text, style_name in reversed(access_content):
        new_p = make_para(doc, text, style_name)
        ref.addprevious(new_p)

doc.save('sprawozdanie.docx')
print('Gotowe!')
