"""
Zastępuje sekcję 10 (Access) sekcją o MongoDB.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

doc = Document('sprawozdanie.docx')

COLOR_H1 = RGBColor(0x1a, 0x1a, 0x2e)
COLOR_H2 = RGBColor(0xe9, 0x45, 0x60)

def get_style_id(doc, style_name):
    for style in doc.styles:
        if style.name == style_name:
            return style.element.get(qn('w:styleId'))
    return 'Normal'

def make_para(doc, text, style_name, color=None):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), get_style_id(doc, style_name))
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    if color:
        rPr = OxmlElement('w:rPr')
        c = OxmlElement('w:color')
        c.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
        rPr.append(c)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

# ── Usuń całą starą sekcję 10 (Access) ────────────────────────────────────
REMOVE_MARKERS = [
    '10. Odwzorowanie schematu w Microsoft Access',
    'Schemat bazy danych został odwzorowany w Microsoft Access',
    'Kluczowe różnice między dialektem PostgreSQL a Access SQL',
    'SERIAL → COUNTER',
    'VARCHAR(n) → TEXT(n)',
    'SMALLINT → SHORT',
    'TIMESTAMPTZ → DATETIME',
    "DEFAULT NOW() → DEFAULT Now()",
    "Ograniczenia CHECK (np. status IN ('scheduled', 'finished')) nie są wspierane w SQL Access",
    '10.1. Kolejność tworzenia tabel',
    'Ze względu na klucze obce, tabele muszą być tworzone w kolejności',
]

for para in list(doc.paragraphs):
    for marker in REMOVE_MARKERS:
        if marker in para.text:
            para._element.getparent().remove(para._element)
            break

# ── Wstaw nową sekcję 10: MongoDB ──────────────────────────────────────────
insert_before = None
for para in doc.paragraphs:
    if para.text.strip() == '11. Testowanie':
        insert_before = para
        break

new_content = [
    ('10. Odwzorowanie schematu w MongoDB',
     'Heading 1', (0x1a, 0x1a, 0x2e)),

    ('Zgodnie z informacjami dodatkowymi do zadania, baza Access może być zastąpiona przez MongoDB. '
     'MongoDB to nierelacyjna baza danych dokumentowa — dane przechowywane są jako dokumenty JSON '
     'w kolekcjach zamiast wierszy w tabelach. Schemat jest elastyczny i nie wymaga z góry '
     'zdefiniowanych kolumn.',
     'Normal', None),

    ('10.1. Porównanie PostgreSQL i MongoDB',
     'Heading 2', (0xe9, 0x45, 0x60)),

    ('Oba systemy przechowują te same dane projektu, ale w różny sposób:',
     'Normal', None),

    ('Struktura: PostgreSQL = tabele + wiersze o stałym schemacie; MongoDB = kolekcje + dokumenty JSON o elastycznej strukturze.',
     'List Bullet', None),

    ('Relacje: PostgreSQL używa kluczy obcych i operacji JOIN; MongoDB osadza powiązane dane bezpośrednio w dokumencie (embedding) lub trzyma referencje id.',
     'List Bullet', None),

    ('Przykład: w PostgreSQL mecz przechowuje home_team_id = 7 i wymaga JOIN z tabelą teams. '
     'W MongoDB dokument meczu zawiera od razu "home_team": "Bayern München" — bez JOINa.',
     'List Bullet', None),

    ('Schemat: PostgreSQL wymaga migracji SQL przy każdej zmianie struktury; MongoDB pozwala dodawać nowe pola do dowolnych dokumentów bez migracji.',
     'List Bullet', None),

    ('Zapytania: PostgreSQL używa SQL (SELECT, JOIN, GROUP BY); MongoDB używa metod find() i potoku agregacji (aggregate).',
     'List Bullet', None),

    ('Transakcje: PostgreSQL obsługuje ACID natywnie; MongoDB wspiera transakcje od wersji 4.0, ale jest mniej naturalny w tym modelu.',
     'List Bullet', None),

    ('10.2. Kolekcje w MongoDB',
     'Heading 2', (0xe9, 0x45, 0x60)),

    ('Baza MongoDB o nazwie typer zawiera 6 kolekcji odpowiadających tabelom PostgreSQL: '
     'scoring_rules, leagues, teams, users, matches, predictions. '
     'Synchronizacja danych odbywa się przez skrypt db/mongo_seed.py, który czyta aktualne dane '
     'z PostgreSQL i wgrywa je do MongoDB — zapewnia spójność obu baz bez ręcznego kopiowania.',
     'Normal', None),

    ('10.3. Schemat kolekcji — MongoDB Compass',
     'Heading 2', (0xe9, 0x45, 0x60)),

    ('Poniżej schemat kolekcji wygenerowany przez MongoDB Compass '
     '(Create → Query Design → Schema → Analyze):',
     'Normal', None),

    ('[ WKLEJ TUTAJ ZRZUT EKRANU SCHEMATU Z MONGODB COMPASS ]',
     'Normal', None),
]

if insert_before:
    ref = insert_before._element
    new_elements = [make_para(doc, text, style, color)
                    for text, style, color in new_content]
    ref.addprevious(new_elements[0])
    for i in range(1, len(new_elements)):
        new_elements[i - 1].addnext(new_elements[i])

doc.save('sprawozdanie.docx')
print('Gotowe!')
