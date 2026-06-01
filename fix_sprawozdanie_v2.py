from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document('sprawozdanie.docx')

# ── 1. Usuń całą błędną sekcję 10 (wstawiona w odwrotnej kolejności) ───────
REMOVE_MARKERS = [
    'Ze względu na klucze obce tabele muszą być tworzone',
    '10.1. Kolejność tworzenia tabel',
    'Ograniczenia CHECK (np. status IN',
    "DEFAULT NOW() → DEFAULT Now()",
    'TIMESTAMPTZ → DATETIME',
    'SMALLINT → SHORT',
    'VARCHAR(n) → TEXT(n)',
    'SERIAL → COUNTER',
    'Kluczowe różnice między dialektem PostgreSQL',
    'Schemat bazy danych został odwzorowany w Microsoft Access',
    '10. Odwzorowanie schematu w Microsoft Access',
]

for para in list(doc.paragraphs):
    for marker in REMOVE_MARKERS:
        if marker in para.text:
            para._element.getparent().remove(para._element)
            break

# ── 2. Poprawki treści istniejących akapitów ───────────────────────────────
for para in doc.paragraphs:
    text = para.text

    # is_admin w opisie tabeli users
    if text.startswith('users: Gracze aplikacji'):
        for run in para.runs:
            if 'hash SHA-256.' in run.text:
                run.text = run.text.replace(
                    'hash SHA-256.',
                    'hash SHA-256. Pole is_admin (domyślnie FALSE) pozwala nadać uprawnienia administratora.'
                )

    # Admin tab — widoczny tylko dla admina
    if text.startswith('Admin — panel wpisywania'):
        for run in para.runs:
            run.text = 'Admin — panel wpisywania wyników meczów z automatycznym przeliczeniem punktów (widoczny wyłącznie dla konta administratora).'

    # Sekcja 9 Migracje — zaktualizuj opis
    if 'Wcześniejsze pliki nigdy nie są modyfikowane' in text:
        for run in para.runs:
            run.text = ''
        para.runs[0].text = (
            'Schemat bazy danych przechowywany jest w dwóch plikach SQL: '
            '001_init.sql zawiera definicje wszystkich tabel, '
            '002_seed.sql wypełnia bazę danymi testowymi. '
            'Do resetowania bazy służy skrypt db/reset_db.sh, który zatrzymuje kontener Docker, '
            'usuwa dane (docker compose down -v), uruchamia go ponownie i wykonuje oba pliki migracji.'
        )

# ── 3. Wstaw sekcję 10 poprawnie (przed "11. Testowanie") ──────────────────
COLOR_H1 = RGBColor(0x1a, 0x1a, 0x2e)
COLOR_H2 = RGBColor(0xe9, 0x45, 0x60)

def get_style_id(doc, style_name):
    for style in doc.styles:
        if style.name == style_name:
            return style.element.get(qn('w:styleId'))
    return 'Normal'

def make_para(doc, text, style_name, color=None):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), get_style_id(doc, style_name))
    pPr.append(pStyle)
    p.append(pPr)
    r = OxmlElement('w:r')
    if color:
        rPr = OxmlElement('w:rPr')
        c = OxmlElement('w:color')
        c.set(qn('w:val'), f'{color[0]:02X}{color[1]:02X}{color[2]:02X}')
        rPr.append(c)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

insert_before = None
for para in doc.paragraphs:
    if para.text.strip() == '11. Testowanie':
        insert_before = para
        break

if insert_before:
    access_content = [
        ('10. Odwzorowanie schematu w Microsoft Access', 'Heading 1', (0x1a, 0x1a, 0x2e)),
        (
            'Schemat bazy danych został odwzorowany w Microsoft Access jako niezależna kopia struktury. '
            'Plik db/access_schema.sql zawiera definicje wszystkich 6 tabel w dialekcie Jet SQL (Access SQL). '
            'Tabele odwzorowują ten sam schemat co baza PostgreSQL — różnią się wyłącznie składnią.',
            'Normal', None
        ),
        ('Kluczowe różnice między dialektem PostgreSQL a Access SQL:', 'Normal', None),
        ('SERIAL → COUNTER (automatyczny licznik klucza głównego)', 'List Bullet', None),
        ('VARCHAR(n) → TEXT(n)', 'List Bullet', None),
        ('SMALLINT → SHORT', 'List Bullet', None),
        ('TIMESTAMPTZ → DATETIME', 'List Bullet', None),
        ("DEFAULT NOW() → DEFAULT Now()", 'List Bullet', None),
        (
            "Ograniczenia CHECK (np. status IN ('scheduled', 'finished')) nie są wspierane "
            'w SQL Access — ustawiane ręcznie przez właściwość Validation Rule w widoku projektu tabeli.',
            'List Bullet', None
        ),
        ('10.1. Kolejność tworzenia tabel', 'Heading 2', (0xe9, 0x45, 0x60)),
        (
            'Ze względu na klucze obce, tabele muszą być tworzone w kolejności: '
            'users, leagues, scoring_rules, następnie teams, matches, a na końcu predictions. '
            'Każdą instrukcję CREATE TABLE należy wykonać osobno w edytorze SQL Accessa '
            '(Create → Query Design → SQL View → Run).',
            'Normal', None
        ),
    ]

    ref = insert_before._element
    # Wstaw pierwszy element przed ref, każdy następny po poprzednim
    new_elements = [make_para(doc, text, style, color) for text, style, color in access_content]
    ref.addprevious(new_elements[0])
    for i in range(1, len(new_elements)):
        new_elements[i - 1].addnext(new_elements[i])

doc.save('sprawozdanie.docx')
print('Gotowe!')
