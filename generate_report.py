"""
Run with: uv run python generate_report.py
Generates sprawozdanie.docx in the project root.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_paragraph(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    return p


def add_placeholder(doc, label):
    """Dashed box for user to fill in (e.g. screenshots)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {label} ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(11)
    # light gray border paragraph via XML shading
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "dashed")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), "AAAAAA")
        pBdr.append(el)
    pPr.append(pBdr)
    doc.add_paragraph()  # spacing after placeholder
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


def add_code(doc, code_text):
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    doc.add_paragraph()


# ─── build document ───────────────────────────────────────────────────────────

doc = Document()

# margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ── STRONA TYTUŁOWA ────────────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("SPRAWOZDANIE Z PROJEKTU")
r.bold = True
r.font.size = Pt(18)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Bazy Danych — projekt zaliczeniowy")
r2.font.size = Pt(13)
r2.italic = True

doc.add_paragraph()

t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Typer Ligowy")
r3.bold = True
r3.font.size = Pt(22)

t4 = doc.add_paragraph()
t4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = t4.add_run("Aplikacja do typowania wyników meczów piłkarskich")
r4.font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

meta = [
    ("Autor:", "[ IMIĘ I NAZWISKO ]"),
    ("Numer indeksu:", "[ NUMER INDEKSU ]"),
    ("Prowadzący:", "[ IMIĘ I NAZWISKO PROWADZĄCEGO ]"),
    ("Rok akademicki:", "2025/2026"),
    ("Data oddania:", "[ DATA ]"),
]
for label, val in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p.add_run(f"{label} ")
    r_l.bold = True
    r_l.font.size = Pt(11)
    r_v = p.add_run(val)
    r_v.font.size = Pt(11)

doc.add_page_break()

# ── 1. OPIS PROJEKTU ───────────────────────────────────────────────────────────
add_heading(doc, "1. Opis projektu", 1)
add_paragraph(doc, (
    "Typer Ligowy to aplikacja webowa umożliwiająca użytkownikom typowanie wyników meczów "
    "piłkarskich rozgrywanych w wybranych ligach (Ekstraklasa, Bundesliga, Liga Mistrzów). "
    "Projekt powstał jako zaliczenie z przedmiotu Bazy Danych i demonstruje projektowanie "
    "schematu relacyjnej bazy danych, pisanie jawnych zapytań SQL oraz integrację "
    "z aplikacją backendową."
))
add_paragraph(doc, (
    "Każdy użytkownik może obstawiać wyniki nadchodzących meczów. Po zakończeniu spotkania "
    "administrator wpisuje rzeczywisty wynik, a system automatycznie przyznaje punkty "
    "wszystkim typującym zgodnie ze skonfigurowanym systemem punktacji. Ranking graczy "
    "jest generowany na żywo jako zapytanie agregujące."
))

add_heading(doc, "1.1. Cel projektu", 2)
doc.add_paragraph("Główne cele dydaktyczne projektu:", style="List Bullet")
doc.add_paragraph("Zaprojektowanie i implementacja schematu bazy danych (8 tabel).", style="List Bullet")
doc.add_paragraph("Pisanie złożonych zapytań SQL (JOIN-y, agregacje, podzapytania).", style="List Bullet")
doc.add_paragraph("Integracja bazy danych z aplikacją backendową (Python + FastAPI).", style="List Bullet")
doc.add_paragraph("Zapewnienie integralności danych przez klucze obce i ograniczenia CHECK.", style="List Bullet")
doc.add_paragraph("Wersjonowanie zmian schematu przez pliki migracji SQL.", style="List Bullet")
doc.add_paragraph()

# ── 2. TECHNOLOGIE ─────────────────────────────────────────────────────────────
add_heading(doc, "2. Użyte technologie", 1)
add_table(doc,
    ["Warstwa", "Technologia", "Uzasadnienie wyboru"],
    [
        ["Baza danych", "PostgreSQL 16", "Dojrzały, relacyjny silnik SQL; wsparcie dla typów timestamptz, smallint, CHECK"],
        ["Backend", "Python 3.11, FastAPI, uvicorn", "Szybkie tworzenie REST API; automatyczny Swagger UI pod /docs"],
        ["Dostęp do DB", "psycopg2-binary", "Jawne SQL bez ORM — zgodnie z wymaganiami projektu"],
        ["Walidacja danych", "Pydantic v2", "Wbudowany w FastAPI; automatyczna walidacja request body"],
        ["Frontend", "HTML + vanilla JS", "Prosta SPA bez zależności; 5 zakładek"],
        ["Testy", "pytest", "Unit testy czystej funkcji punktującej"],
        ["Środowisko", "Docker, uv", "Powtarzalne środowisko; izolacja bazy w kontenerze"],
    ]
)

# ── 3. SCHEMAT BAZY DANYCH ─────────────────────────────────────────────────────
add_heading(doc, "3. Schemat bazy danych", 1)
add_paragraph(doc, (
    "Baza danych zawiera 8 tabel. Schemat został zaprojektowany tak, aby unikać redundancji "
    "danych i zachować spójność referencyjną przez klucze obce."
))

add_heading(doc, "3.1. Diagram ER", 2)
add_placeholder(doc, "WKLEJ TUTAJ ZRZUT DIAGRAMU ER Z DBDIAGRAM.IO")

add_heading(doc, "3.2. Opis tabel", 2)

tables_desc = [
    ("users", "Gracze aplikacji. Unikalne pola: nick, email. Hasło przechowywane jako hash SHA-256."),
    ("leagues", "Rozgrywki ligowe (Ekstraklasa, Bundesliga, Liga Mistrzów). Pole season identyfikuje edycję (np. 2024/25)."),
    ("teams", "Drużyny przypisane do ligi przez klucz obcy league_id."),
    ("matches", "Mecze z datą kickoff_at, identyfikatorami drużyn i pola wynikowe (home_goals, away_goals). Status: scheduled → finished."),
    ("scoring_rules", "Konfiguracja systemu punktacji (exact_pts=5, diff_pts=3, tendency_pts=2). Przechowywana w bazie zamiast hardkodowania."),
    ("predictions", "Typy użytkowników. Unikalne ograniczenie (user_id, match_id) — jeden typ na mecz. Pole points_awarded uzupełniane po zakończeniu meczu."),
    ("private_leagues", "Prywatne ligi z unikalnym kodem dołączenia. Każda liga ma własny zestaw reguł punktacji."),
    ("private_league_members", "Tabela łącząca M:N między użytkownikami a prywatnymi ligami."),
]

for name, desc in tables_desc:
    p = doc.add_paragraph(style="List Bullet")
    r_name = p.add_run(f"{name}: ")
    r_name.bold = True
    r_name.font.name = "Courier New"
    p.add_run(desc)

doc.add_paragraph()

add_heading(doc, "3.3. Kluczowe ograniczenia i indeksy", 2)
add_paragraph(doc, "Poniżej najważniejsze ograniczenia zaimplementowane w schemacie:")
doc.add_paragraph("UNIQUE(user_id, match_id) w tabeli predictions — jeden typ na mecz.", style="List Bullet")
doc.add_paragraph("CHECK (home_team_id <> away_team_id) w tabeli matches — drużyna nie może grać sama ze sobą.", style="List Bullet")
doc.add_paragraph("CHECK (status IN ('scheduled', 'finished')) — tylko dwa dozwolone statusy meczu.", style="List Bullet")
doc.add_paragraph("UNIQUE(join_code) w tabeli private_leagues — unikalny kod dołączenia.", style="List Bullet")
doc.add_paragraph("NOT NULL na wszystkich wymaganych polach; klucze obce z ON DELETE CASCADE tam, gdzie logicznie uzasadnione.", style="List Bullet")
doc.add_paragraph()

# ── 4. ARCHITEKTURA APLIKACJI ──────────────────────────────────────────────────
add_heading(doc, "4. Architektura aplikacji", 1)
add_paragraph(doc, (
    "Aplikacja zbudowana jest w architekturze warstwowej. Każda warstwa ma jasno określoną "
    "odpowiedzialność i nie importuje kodu z wyższej lub tej samej warstwy w niedozwolony sposób."
))

add_code(doc,
    "HTTP Client (przeglądarka / Swagger UI)\n"
    "   ↓\n"
    "REST API (FastAPI, routers/)     ← przyjmuje request, zwraca response\n"
    "   ↓\n"
    "Logika domenowa (domain/)        ← czysta, bez I/O, testowalna\n"
    "   ↓\n"
    "Dostęp do danych (db/)           ← parametryzowane SQL-e\n"
    "   ↓\n"
    "PostgreSQL 16"
)

add_heading(doc, "4.1. Warstwa domenowa (domain/)", 2)
add_paragraph(doc, (
    "Najważniejsza część projektu pod kątem dydaktycznym. Funkcje w tej warstwie są czyste "
    "(pure functions) — nie wykonują żadnych operacji I/O, nie importują FastAPI ani psycopg2. "
    "Dzięki temu są w pełni testowalne jednostkowo."
))
doc.add_paragraph("scoring.py — funkcja calculate_points() obliczająca punkty dla jednego typu.", style="List Bullet")
doc.add_paragraph("match_results.py — przelicza punkty dla wszystkich typów danego meczu.", style="List Bullet")
doc.add_paragraph("predictions.py — waliduje czy typ można jeszcze dodać (kickoff_at > teraz).", style="List Bullet")
doc.add_paragraph()

add_heading(doc, "4.2. Warstwa dostępu do danych (db/)", 2)
add_paragraph(doc, (
    "Wszystkie zapytania SQL przechowywane są jako stałe stringi w db/queries.py. "
    "Parametry przekazywane przez %s — nigdy f-stringi (ochrona przed SQL Injection). "
    "Pula połączeń zarządzana przez psycopg2.pool.SimpleConnectionPool."
))

add_heading(doc, "4.3. Warstwa API (routers/)", 2)
add_paragraph(doc, (
    "Routery odpowiadają wyłącznie za obsługę HTTP: odbiór request body (walidowany przez Pydantic), "
    "wywołanie logiki domenowej lub funkcji z db/, i zwrócenie odpowiedzi. "
    "Routery nie wykonują SQL bezpośrednio."
))

# ── 5. SYSTEM PUNKTACJI ────────────────────────────────────────────────────────
add_heading(doc, "5. System punktacji", 1)
add_paragraph(doc, (
    "Liczy się tylko najwyższy próg trafienia — punkty się nie sumują. "
    "Wartości punktów są konfigurowalne i przechowywane w tabeli scoring_rules."
))
add_table(doc,
    ["Trafienie", "Warunek", "Punkty"],
    [
        ["Dokładny wynik", "pred_home == real_home AND pred_away == real_away", "5"],
        ["Trafiona różnica bramek", "pred_home − pred_away == real_home − real_away", "3"],
        ["Trafiony rezultat (kto wygrał)", "sign(pred_home − pred_away) == sign(real_home − real_away)", "2"],
        ["Pudło", "żaden z powyższych warunków", "0"],
    ]
)
add_paragraph(doc, (
    "Uwaga: trafiony remis (np. 1:1 → 2:2) zawsze daje 3 pkt, nigdy 2. "
    "Wynika to z warunku różnicy bramek: 0 = 0, więc remis trafia na wyższy próg."
), italic=True)

add_heading(doc, "5.1. Implementacja funkcji punktującej", 2)
add_code(doc,
    "def calculate_points(pred_home, pred_away, real_home, real_away, rules) -> int:\n"
    "    if pred_home == real_home and pred_away == real_away:\n"
    "        return rules['exact_pts']\n"
    "    if (pred_home - pred_away) == (real_home - real_away):\n"
    "        return rules['diff_pts']\n"
    "    pred_sign = (pred_home - pred_away > 0) - (pred_home - pred_away < 0)\n"
    "    real_sign = (real_home - real_away > 0) - (real_home - real_away < 0)\n"
    "    if pred_sign == real_sign:\n"
    "        return rules['tendency_pts']\n"
    "    return 0"
)

# ── 6. ENDPOINTY API ───────────────────────────────────────────────────────────
add_heading(doc, "6. Endpointy API", 1)
add_table(doc,
    ["Metoda", "Ścieżka", "Opis", "Kod sukcesu"],
    [
        ["POST", "/auth/login", "Logowanie — weryfikacja hasła SHA-256", "200"],
        ["POST", "/auth/register", "Rejestracja nowego użytkownika", "201"],
        ["GET", "/matches", "Lista wszystkich meczów z JOIN-ami do drużyn i ligi", "200"],
        ["GET", "/matches/{id}", "Szczegóły jednego meczu", "200"],
        ["POST", "/predictions", "Dodaj typ (walidacja: mecz nie może być started)", "201"],
        ["GET", "/predictions/user/{id}", "Historia typów użytkownika z wynikami", "200"],
        ["GET", "/ranking", "Ranking — agregacja punktów z finished meczów", "200"],
        ["POST", "/matches/{id}/result", "Admin: wpisz wynik + przelicz punkty atomowo", "200"],
    ]
)

add_heading(doc, "6.1. Przykładowe odpowiedzi API", 2)
add_paragraph(doc, "GET /matches — fragment odpowiedzi:")
add_code(doc,
    '[\n'
    '  {\n'
    '    "id": 1,\n'
    '    "home_team": "Lechia Gdańsk",\n'
    '    "away_team": "Lech Poznań",\n'
    '    "league": "Ekstraklasa",\n'
    '    "kickoff_at": "2026-05-25T17:00:00+00:00",\n'
    '    "home_goals": 2,\n'
    '    "away_goals": 1,\n'
    '    "status": "finished"\n'
    '  }\n'
    ']'
)

# ── 7. FRONTEND ────────────────────────────────────────────────────────────────
add_heading(doc, "7. Interfejs graficzny", 1)
add_paragraph(doc, (
    "Frontend to pojedynczy plik HTML (SPA — Single Page Application) oparty na vanilla JavaScript "
    "bez zewnętrznych zależności. Komunikuje się z API przez fetch(). Sesja przechowywana "
    "w localStorage jako obiekt {id, nick}."
))
add_paragraph(doc, "Aplikacja zawiera 5 zakładek:")
doc.add_paragraph("Mecze — lista wszystkich meczów z statusem i wynikami.", style="List Bullet")
doc.add_paragraph("Dodaj typ — formularz typowania dla zaplanowanych meczów.", style="List Bullet")
doc.add_paragraph("Moje typy — historia typów zalogowanego użytkownika z kolorowymi wskaźnikami punktów.", style="List Bullet")
doc.add_paragraph("Ranking — tabela liderów aktualizowana na żywo.", style="List Bullet")
doc.add_paragraph("Admin — panel wpisywania wyników meczów z automatycznym przeliczeniem punktów.", style="List Bullet")
doc.add_paragraph()

add_heading(doc, "7.1. Zrzuty ekranu", 2)
add_placeholder(doc, "ZRZUT EKRANU — EKRAN LOGOWANIA")
add_placeholder(doc, "ZRZUT EKRANU — ZAKŁADKA MECZE")
add_placeholder(doc, "ZRZUT EKRANU — ZAKŁADKA DODAJ TYP")
add_placeholder(doc, "ZRZUT EKRANU — ZAKŁADKA MOJE TYPY (z punktami)")
add_placeholder(doc, "ZRZUT EKRANU — ZAKŁADKA RANKING")
add_placeholder(doc, "ZRZUT EKRANU — PANEL ADMINA")
add_placeholder(doc, "ZRZUT EKRANU — SWAGGER UI (/docs)")

# ── 8. PRZYKŁADOWE ZAPYTANIA SQL ───────────────────────────────────────────────
add_heading(doc, "8. Kluczowe zapytania SQL", 1)

add_heading(doc, "8.1. Ranking graczy", 2)
add_code(doc,
    "SELECT\n"
    "    u.id   AS user_id,\n"
    "    u.nick,\n"
    "    COALESCE(SUM(p.points_awarded), 0) AS total_points\n"
    "FROM users u\n"
    "LEFT JOIN predictions p ON p.user_id = u.id\n"
    "LEFT JOIN matches m     ON p.match_id = m.id AND m.status = 'finished'\n"
    "GROUP BY u.id, u.nick\n"
    "ORDER BY total_points DESC;"
)
add_paragraph(doc, (
    "Zapytanie używa LEFT JOIN, żeby gracze bez żadnych typów pojawiali się w rankingu z 0 punktami. "
    "COALESCE obsługuje przypadek NULL (brak typów). Warunek m.status = 'finished' "
    "w klauzuli JOIN (nie WHERE) zapewnia, że gracze bez zakończonych meczów nadal trafiają do wyników."
))

add_heading(doc, "8.2. Historia typów użytkownika", 2)
add_code(doc,
    "SELECT\n"
    "    p.id, p.pred_home, p.pred_away, p.points_awarded, p.created_at,\n"
    "    m.kickoff_at, m.status, m.home_goals, m.away_goals,\n"
    "    ht.name AS home_team,\n"
    "    at.name AS away_team,\n"
    "    l.name  AS league\n"
    "FROM predictions p\n"
    "JOIN matches m  ON p.match_id = m.id\n"
    "JOIN teams ht   ON m.home_team_id = ht.id\n"
    "JOIN teams at   ON m.away_team_id = at.id\n"
    "JOIN leagues l  ON m.league_id = l.id\n"
    "WHERE p.user_id = %s\n"
    "ORDER BY m.kickoff_at DESC;"
)

add_heading(doc, "8.3. Transakcja: wynik meczu + przeliczenie punktów", 2)
add_paragraph(doc, (
    "Kluczowa operacja w systemie. Wszystkie UPDATE punktów i zmiana statusu meczu "
    "wykonywane są w jednej transakcji — albo wszystko się uda, albo nic."
))
add_code(doc,
    "-- krok 1: zaktualizuj punkty dla każdego typu (N razy)\n"
    "UPDATE predictions SET points_awarded = %s WHERE id = %s;\n\n"
    "-- krok 2: ustaw wynik i zmień status na finished (jeden commit)\n"
    "UPDATE matches\n"
    "SET home_goals = %s, away_goals = %s, status = 'finished'\n"
    "WHERE id = %s;\n\n"
    "COMMIT;"
)

# ── 9. MIGRACJE SCHEMATU ───────────────────────────────────────────────────────
add_heading(doc, "9. Migracje schematu", 1)
add_paragraph(doc, (
    "Zmiany schematu bazy danych wersjonowane są przez kolejne pliki SQL. "
    "Wcześniejsze pliki nigdy nie są modyfikowane — każda zmiana to nowy plik."
))
add_table(doc,
    ["Plik migracji", "Zawartość"],
    [
        ["001_init.sql", "Tworzenie wszystkich 8 tabel ze wszystkimi ograniczeniami"],
        ["002_seed.sql", "Pierwsze dane testowe (zastąpione przez 004)"],
        ["003_add_auth.sql", "Dodanie kolumny password_hash do tabeli users (ALTER TABLE)"],
        ["004_reseed.sql", "Pełne przeładowanie danych: 4 ligi, 13 drużyn, 9 graczy, 8 meczów"],
        ["005_fix_ranking.sql", "Korekta wartości points_awarded dla uzyskania docelowego rankingu"],
    ]
)

# ── 10. TESTOWANIE ─────────────────────────────────────────────────────────────
add_heading(doc, "10. Testowanie", 1)
add_paragraph(doc, (
    "Funkcja punktująca pokryta jest 8 testami jednostkowymi obejmującymi wszystkie "
    "przypadki brzegowe systemu punktacji."
))
add_table(doc,
    ["Typ", "Rzeczywisty", "Oczekiwane punkty", "Opis"],
    [
        ["2:1", "2:1", "5", "Dokładny wynik"],
        ["2:1", "3:2", "3", "Trafiona różnica (+1)"],
        ["2:1", "1:0", "3", "Trafiona różnica (+1)"],
        ["2:1", "4:0", "2", "Tylko rezultat (wygrana)"],
        ["1:1", "2:2", "3", "Remis → różnica (0=0)"],
        ["1:1", "0:0", "3", "Remis → różnica (0=0)"],
        ["2:1", "0:2", "0", "Pudło (porażka zamiast wygranej)"],
        ["0:0", "1:0", "0", "Pudło (wygrana zamiast remisu)"],
    ]
)

add_heading(doc, "10.1. Wyniki testów", 2)
add_placeholder(doc, "ZRZUT EKRANU — uv run pytest (wszystkie testy zielone)")

# ── 11. WNIOSKI ────────────────────────────────────────────────────────────────
add_heading(doc, "11. Wnioski", 1)
add_placeholder(doc, "WYPEŁNIJ: wnioski z projektu — co zadziałało, co było trudne, czego się nauczyłeś")

# ── LITERATURA ─────────────────────────────────────────────────────────────────
add_heading(doc, "Literatura i zasoby", 1)
doc.add_paragraph("Dokumentacja PostgreSQL 16 — https://www.postgresql.org/docs/16/", style="List Number")
doc.add_paragraph("Dokumentacja FastAPI — https://fastapi.tiangolo.com/", style="List Number")
doc.add_paragraph("Dokumentacja psycopg2 — https://www.psycopg.org/docs/", style="List Number")
doc.add_paragraph("Pydantic v2 docs — https://docs.pydantic.dev/", style="List Number")
doc.add_paragraph("dbdiagram.io — narzędzie do tworzenia diagramów ER", style="List Number")

# ─── save ──────────────────────────────────────────────────────────────────────
doc.save("sprawozdanie.docx")
print("Wygenerowano: sprawozdanie.docx")
